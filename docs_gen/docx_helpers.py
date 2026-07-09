"""Reusable helpers for building the Technical Design Document with python-docx.

Kept separate from the content script (build_doc.py) so the ~31-chapter content generation
stays readable - this module only knows how to lay out styled building blocks (headings,
tables, code blocks, TOC/page-number fields), never what to put in them.
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x2E, 0x74, 0xB5)
DARK_GRAY = RGBColor(0x40, 0x40, 0x40)
LIGHT_GRAY_FILL = "F2F2F2"
CODE_FILL = "2B2B2B"
CODE_FG = RGBColor(0xE0, 0xE0, 0xE0)


def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def build_base_document():
    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for level, size, color, space_before in (
        (1, 20, NAVY, 24),
        (2, 15, NAVY, 18),
        (3, 12.5, ACCENT, 12),
        (4, 11, DARK_GRAY, 8),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(8 if level == 1 else 6)
        style.paragraph_format.keep_with_next = True

    # Title/Subtitle styles used only on the cover page.
    title_style = doc.styles.add_style("CoverTitle", 1)
    title_style.base_style = doc.styles["Title"]
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(34)
    title_style.font.bold = True
    title_style.font.color.rgb = NAVY

    subtitle_style = doc.styles.add_style("CoverSubtitle", 1)
    subtitle_style.base_style = doc.styles["Normal"]
    subtitle_style.font.name = "Calibri Light"
    subtitle_style.font.size = Pt(16)
    subtitle_style.font.color.rgb = ACCENT

    code_style = doc.styles.add_style("CodeBlock", 1)
    code_style.base_style = doc.styles["Normal"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = CODE_FG
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(10)
    code_style.paragraph_format.left_indent = Cm(0.3)

    caption_style = doc.styles.add_style("FigureCaption", 1)
    caption_style.base_style = doc.styles["Normal"]
    caption_style.font.name = "Calibri"
    caption_style.font.size = Pt(9)
    caption_style.font.italic = True
    caption_style.font.color.rgb = DARK_GRAY

    note_style = doc.styles.add_style("NoteBox", 1)
    note_style.base_style = doc.styles["Normal"]
    note_style.font.size = Pt(10)
    note_style.font.italic = True

    _add_header_footer(doc)
    return doc


def _add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = "Intelligent Credit Card Rewards Agent — Technical Design Document"
    header_p.style = doc.styles["Normal"]
    for run in header_p.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = DARK_GRAY
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_field(footer_p, "PAGE")
    run = footer_p.add_run(" of ")
    run.font.size = Pt(9)
    _add_field(footer_p, "NUMPAGES")
    for run in footer_p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = DARK_GRAY


def _add_field(paragraph, field_code):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    return run


def add_cover_page(doc, title, subtitle, meta_rows, revision_rows):
    p = doc.add_paragraph(style="CoverTitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)

    sp = doc.add_paragraph(style="CoverSubtitle")
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.add_run(subtitle)

    for _ in range(3):
        doc.add_paragraph()

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, value in meta_rows:
        row = table.add_row()
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(3.6)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)

    doc.add_paragraph()
    h = doc.add_heading("Revision History", level=3)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table(doc, ["Version", "Date", "Author", "Description"], revision_rows)
    doc.add_page_break()


def add_toc_page(doc):
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "Right-click here and choose “Update Field” to generate the Table of Contents."
    )
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(placeholder)
    run._r.append(fld_char_end)
    doc.add_page_break()


def add_table(doc, headers, rows, widths=None, header_fill="1F3A5F"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
        set_cell_background(hdr_cells[i], header_fill)
    for row_data in rows:
        cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(cell_text))
            run.font.size = Pt(9.5)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_code_block(doc, code_text):
    p = doc.add_paragraph(style="CodeBlock")
    lines = code_text.strip("\n").split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line if i == 0 else "\n" + line)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1E1E1E")
    p_pr.append(shd)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
    return p


def add_diagram_block(doc, diagram_text):
    """ASCII/Unicode box diagrams - monospace, boxed, but light-background (distinct from
    dark code blocks) so it reads as a diagram, not source code."""
    p = doc.add_paragraph(style="CodeBlock")
    lines = diagram_text.strip("\n").split("\n")
    for i, line in enumerate(lines):
        p.add_run(line if i == 0 else "\n" + line)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    p_pr.append(shd)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
        run.font.size = Pt(8.5)
    return p


def add_note(doc, label, text, color=RGBColor(0x2E, 0x74, 0xB5)):
    p = doc.add_paragraph(style="NoteBox")
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = color
    p.add_run(text)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EEF3FA")
    p_pr.append(shd)
    return p


def add_bullets(doc, items, style="List Bullet"):
    for item in items:
        doc.add_paragraph(item, style=style)


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")
